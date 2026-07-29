from django.shortcuts import render
from tasks.models import tasks
from django.http import HttpResponse

from openpyxl import Workbook
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors


def task_report(request):

    # ---------------- GET ALL TASKS ----------------
    task_list = tasks.objects.all()

    # ---------------- SEARCH ----------------
    search = request.GET.get("search")
    if search:
        task_list = task_list.filter(title__icontains=search)

    # ---------------- STATUS FILTER ----------------
    status = request.GET.get("status")
    if status and status != "all":
        task_list = task_list.filter(status=status)

    # ---------------- DATE FILTER ----------------
    date = request.GET.get("date")
    if date:
        task_list = task_list.filter(created_at__date=date)

    # ---------------- EXPORT ----------------
    export = request.GET.get("export")

    # ==================================================
    # EXCEL EXPORT
    # ==================================================
    if export == "excel":

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        response["Content-Disposition"] = 'attachment; filename="Task_Report.xlsx"'

        wb = Workbook()
        ws = wb.active
        ws.title = "Task Report"

        ws.append([
            "Task Title",
            "Project",
            "Employee",
            "Priority",
            "Status",
            "Due Date",
            "Created At"
        ])

        for task in task_list:
            ws.append([
                task.title,
                str(task.project),
                str(task.employee),
                task.priority,
                task.status,
                task.due_date.strftime("%Y-%m-%d"),
                task.created_at.strftime("%Y-%m-%d %H:%M"),
            ])

        wb.save(response)
        return response

    # ==================================================
    # PDF EXPORT
    # ==================================================
    elif export == "pdf":

        response = HttpResponse(content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="Task_Report.pdf"'

        document = SimpleDocTemplate(response)

        data = [[
            "Title",
            "Project",
            "Employee",
            "Priority",
            "Status",
            "Due Date"
        ]]

        for task in task_list:
            data.append([
                task.title,
                str(task.project),
                str(task.employee),
                task.priority,
                task.status,
                str(task.due_date),
            ])

        table = Table(data)

        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 10),
        ]))

        document.build([table])

        return response

    # ==================================================
    # WORD EXPORT
    # ==================================================
    elif export == "word":

        document = Document()

        document.add_heading("Task Report", level=1)

        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Title"
        header[1].text = "Project"
        header[2].text = "Employee"
        header[3].text = "Priority"
        header[4].text = "Status"
        header[5].text = "Due Date"

        for task in task_list:

            row = table.add_row().cells

            row[0].text = task.title
            row[1].text = str(task.project)
            row[2].text = str(task.employee)
            row[3].text = task.priority
            row[4].text = task.status
            row[5].text = str(task.due_date)

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        response["Content-Disposition"] = 'attachment; filename="Task_Report.docx"'

        document.save(response)

        return response

    # ---------------- REPORT COUNTS ----------------

    completed_tasks = task_list.filter(status="Completed").count()
    progress_tasks = task_list.filter(status="In Progress").count()
    pending_tasks = task_list.filter(status="Pending").count()
    review_tasks = task_list.filter(status="Review").count()
    total_tasks = task_list.count()

    context = {
        "completed_tasks": completed_tasks,
        "progress_tasks": progress_tasks,
        "pending_tasks": pending_tasks,
        "review_tasks": review_tasks,
        "total_tasks": total_tasks,
        "tasks": task_list,
        "search": search,
        "status": status,
        "date": date,
    }

    return render(request, "reports/task_report.html", context)